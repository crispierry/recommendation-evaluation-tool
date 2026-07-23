#!/usr/bin/env python3
"""Build the public synthetic executive report from the frozen analysis JSON."""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INK = "1A1A1A"
MUTED = "5F6368"
ACCENT = "96683D"
LIGHT = "F6F1EA"
SOFT = "F2F4F7"
WHITE = "FFFFFF"
PAGE_WIDTH_DXA = 9360


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--analysis", required=True)
    parser.add_argument("--output", required=True)
    parser.add_argument("--chart-dir", required=True)
    args = parser.parse_args()

    analysis_path = Path(args.analysis)
    output_path = Path(args.output)
    chart_dir = Path(args.chart_dir)
    chart_dir.mkdir(parents=True, exist_ok=True)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    analysis = json.loads(analysis_path.read_text(encoding="utf-8"))
    if analysis.get("synthetic") is not True:
        raise ValueError("The public report accepts only explicitly synthetic analysis.")

    routine_chart = chart_dir / "routine-overlap.png"
    repetition_chart = chart_dir / "repetition-recurrence.png"
    build_routine_chart(analysis, routine_chart)
    build_repetition_chart(analysis, repetition_chart)
    document = build_document(analysis, routine_chart, repetition_chart)
    document.save(output_path)
    print(f"Built synthetic Word report: {output_path}")


def build_document(analysis: dict, routine_chart: Path, repetition_chart: Path) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    configure_styles(doc)
    configure_properties(doc)
    configure_header_footer(section)
    add_cover(doc, analysis)
    doc.add_page_break()
    add_executive_summary(doc, analysis)
    add_metric_strip(doc, analysis)
    add_heading(doc, "Two questions, two analytical grains", 1)
    add_body(
        doc,
        "Day-over-day continuity asks whether a title in the current profile's first 50 positions appeared "
        "anywhere in the prior synthetic day's first 50. Exact-clip recurrence asks whether the same "
        "canonical clip appeared in an earlier controlled repetition run. The measures illuminate different "
        "risks and should never be combined into one score.",
    )
    add_heading(doc, "Five synthetic days reveal distinct continuity patterns", 1)
    add_body(
        doc,
        "The three routine profiles use the same catalog and capture window but different configured behavior "
        "bands. That makes the evaluation sensitive to profile-level delivery patterns without implying that "
        "the synthetic strategies describe real people or a production ranker.",
    )
    add_figure(doc, routine_chart, "Figure 1. Title overlap with the immediately prior synthetic day.")
    add_routine_table(doc, analysis)

    add_heading(doc, "Controlled runs isolate exact-clip recurrence", 1)
    add_body(
        doc,
        "The repetition study uses three different profiles and only 20 positions per run. A repeat is counted "
        "only when the same canonical clip ID appeared in an earlier run for that profile. Another clip from "
        "the same fictional title remains fresh.",
    )
    add_figure(doc, repetition_chart, "Figure 2. Progressive exact-clip recurrence in Runs 2 and 3.")
    add_repetition_table(doc, analysis)

    add_heading(doc, "What the evaluation system makes visible", 1)
    for text in [
        "Profile-level differences that an aggregate freshness number would hide.",
        "The distinction between repeated titles and repeated exact clips.",
        "Catalog exposure, concentration, and the number of unique clips actually observed.",
        "A small human-review queue for ambiguous identity, possible duplicate scenes, and incomplete metadata.",
        "A reproducible record of configuration, denominators, exclusions, and computed findings.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "Configuration and reproducibility", 1)
    add_body(
        doc,
        "The published run is deterministic. Teams can change the JSON configuration or validated command-line "
        "overrides to test different catalog sizes, clip ranges, seeds, profile counts, days, positions per day, "
        "repetition runs, and recurrence strategies. Changing the configuration regenerates the dataset, artwork, "
        "dashboard metrics, and this report from the same source.",
    )
    add_config_table(doc, analysis)

    add_heading(doc, "Interpretation limits", 1)
    for limitation in analysis["limitations"]:
        add_bullet(doc, limitation)

    add_heading(doc, "Recommended next evaluations", 1)
    for text in [
        "Add position-weighted measures to distinguish top-of-feed repetition from lower-position recurrence.",
        "Run longer synthetic windows to test whether profile ordering remains stable as the anchor day moves.",
        "Introduce controlled catalog shocks to evaluate how quickly the system detects reduced novelty.",
        "Track review precision against the synthetic ground truth without allowing visitor decisions to alter the frozen report.",
    ]:
        add_numbered(doc, text)

    add_body(
        doc,
        "This document and every visual in it were generated from fictional titles, fictional profiles, "
        "fictional artwork, and deterministic synthetic recommendation runs.",
        color=ACCENT,
        bold=True,
        before=12,
    )
    return doc


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Title", 28, INK, 0, 8),
        ("Subtitle", 13, MUTED, 0, 16),
        ("Heading 1", 17, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 11, INK, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        p_borders = style._element.pPr.find(qn("w:pBdr"))
        if p_borders is not None:
            style._element.pPr.remove(p_borders)

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def configure_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "Short-Form Recommendation Evaluation Lab"
    props.subject = "Deterministic synthetic evaluation of catalog continuity and exact-clip recurrence"
    props.author = "Cristiano Pierry"
    props.last_modified_by = "Cristiano Pierry"
    props.comments = "Generated exclusively from synthetic public data."
    props.category = "Synthetic product evaluation"
    props.keywords = "recommendation evaluation, synthetic data, feed quality, repetition"


def configure_header_footer(section) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("RECOMMENDATION EVALUATION LAB  ·  SYNTHETIC")
    set_run(run, 8, MUTED, bold=True)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, 8, MUTED)
    add_field(paragraph, "PAGE")


def add_cover(doc: Document, analysis: dict) -> None:
    add_spacer(doc, 54)
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(18)
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("EXECUTIVE READOUT  ·  SYNTHETIC DATA")
    set_run(run, 10, ACCENT, bold=True)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Short-Form Recommendation\nEvaluation Lab")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Catalog continuity, exact-clip recurrence, and human review")

    add_spacer(doc, 26)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3120, 3120, 3120], header=False)
    values = [
        ("500", "fictional titles"),
        ("930", "synthetic appearances"),
        ("5 + 3", "days and controlled runs"),
    ]
    for cell, (value, label) in zip(table.rows[0].cells, values):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        value_run = p.add_run(value)
        set_run(value_run, 20, INK, bold=True)
        p.add_run("\n")
        label_run = p.add_run(label.upper())
        set_run(label_run, 8, MUTED, bold=True)

    add_spacer(doc, 34)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run(
        "Every title, profile, recommendation, metric, poster, and screen capture in this report is fictional."
    )
    set_run(run, 10, MUTED, italic=True)


def add_executive_summary(doc: Document, analysis: dict) -> None:
    add_heading(doc, "Executive Summary", 1)
    routine = analysis["routineProfiles"]
    repetition = analysis["repetitionProfiles"]
    highest_routine = max(routine, key=lambda item: item["meanRate"])
    lowest_routine = min(routine, key=lambda item: item["meanRate"])
    highest_repeat = max(repetition, key=lambda item: item["latestRate"])
    lowest_repeat = min(repetition, key=lambda item: item["latestRate"])

    bullets = [
        (
            "The evaluation detects meaningful profile differences. "
            f"Mean day-over-day title overlap ranges from {pct(lowest_routine['meanRate'])} "
            f"for {lowest_routine['profileLabel']} to {pct(highest_routine['meanRate'])} "
            f"for {highest_routine['profileLabel']}, a {analysis['headline']['routineSpreadPoints']:.0f}-point spread."
        ),
        (
            "The controlled repetition study separates low and high exact-clip recurrence. "
            f"Run 3 ranges from {pct(lowest_repeat['latestRate'])} for {lowest_repeat['profileLabel']} "
            f"to {pct(highest_repeat['latestRate'])} for {highest_repeat['profileLabel']}."
        ),
        (
            f"The 930 appearances expose {analysis['counts']['exposedTitles']} of 500 fictional titles "
            f"and {analysis['counts']['exposedClips']} of 2,000 canonical clips. Coverage is a property "
            "of the configured sample, not a production benchmark."
        ),
        (
            "The result is an evaluation demonstration, not a causal diagnosis. It shows what the "
            "configured synthetic delivery patterns look like when measured consistently."
        ),
    ]
    for text in bullets:
        add_bullet(doc, text)


def add_metric_strip(doc: Document, analysis: dict) -> None:
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2340, 2340, 2340, 2340], header=False)
    metrics = [
        (f"{analysis['headline']['routineSpreadPoints']:.0f} pts", "routine spread"),
        (f"{analysis['headline']['repetitionSpreadPoints']:.0f} pts", "recurrence spread"),
        (pct(analysis["headline"]["titleCoverageRate"]), "title coverage"),
        (str(analysis["counts"]["reviewCases"]), "review cases"),
    ]
    for cell, (value, label) in zip(table.rows[0].cells, metrics):
        shade_cell(cell, SOFT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(value)
        set_run(run, 16, INK, bold=True)
        p.add_run("\n")
        run = p.add_run(label.upper())
        set_run(run, 7.5, MUTED, bold=True)


def add_routine_table(doc: Document, analysis: dict) -> None:
    table = doc.add_table(rows=1, cols=4)
    headers = ["Profile", "Strategy", "Mean overlap", "Day 4 → Day 5"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for profile in analysis["routineProfiles"]:
        row = table.add_row().cells
        row[0].text = profile["profileLabel"]
        row[1].text = profile["strategy"]
        row[2].text = pct(profile["meanRate"])
        row[3].text = pct(profile["latestRate"])
    set_table_geometry(table, [3000, 2600, 1880, 1880], header=True)


def add_repetition_table(doc: Document, analysis: dict) -> None:
    table = doc.add_table(rows=1, cols=5)
    headers = ["Profile", "Strategy", "Run 2", "Run 3", "Unique clips"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for profile in analysis["repetitionProfiles"]:
        row = table.add_row().cells
        row[0].text = profile["profileLabel"]
        row[1].text = profile["strategy"]
        row[2].text = pct(profile["progressive"][1]["rate"])
        row[3].text = pct(profile["progressive"][2]["rate"])
        row[4].text = str(profile["cumulativeUniqueClips"])
    set_table_geometry(table, [2800, 2200, 1400, 1400, 1560], header=True)


def add_config_table(doc: Document, analysis: dict) -> None:
    table = doc.add_table(rows=1, cols=3)
    headers = ["Setting", "Published value", "What changes"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    rows = [
        ("Catalog", "500 titles · 3–5 clips each", "Size and clip diversity"),
        ("Routine run", "3 profiles · 5 days · 50 clips", "Continuity window and sample"),
        ("Repetition run", "3 profiles · 3 runs · 20 clips", "Controlled recurrence window"),
        ("Behavior", "Distinct configured profile bands", "Novelty and recurrence mix"),
        ("Seed", analysis["configDigest"][:12], "Complete deterministic replay"),
    ]
    for values in rows:
        row = table.add_row().cells
        for cell, value in zip(row, values):
            cell.text = value
    set_table_geometry(table, [2100, 3420, 3840], header=True)


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_body(
    doc: Document,
    text: str,
    *,
    color: str = INK,
    bold: bool = False,
    italic: bool = False,
    before: int = 0,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    run = paragraph.add_run(text)
    set_run(run, 10.5, color, bold=bold, italic=italic)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def add_numbered(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.add_run(text)


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    picture = run.add_picture(str(image_path), width=Inches(6.1))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", caption.split(".")[0])
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(4)
    caption_paragraph.paragraph_format.space_after = Pt(8)
    caption_paragraph.paragraph_format.keep_with_next = True
    run = caption_paragraph.add_run(caption)
    set_run(run, 8.5, MUTED, italic=True)


def add_spacer(doc: Document, points: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(points)


def set_run(run, size: float, color: str, *, bold=False, italic=False) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def set_table_geometry(table, widths: list[int], *, header: bool) -> None:
    if sum(widths) != PAGE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {PAGE_WIDTH_DXA}: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            row._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
        for index, (cell, width) in enumerate(zip(row.cells, widths)):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell, 90, 90, 120, 120)
            if header and row_index == 0:
                shade_cell(cell, SOFT)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                if index >= 2:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run(run, 9, INK, bold=header and row_index == 0)


def set_cell_margins(cell, top: int, bottom: int, start: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run(run, 8, MUTED)


def build_routine_chart(analysis: dict, destination: Path) -> None:
    labels = [item["profileLabel"].split(" — ")[-1] for item in analysis["routineProfiles"]]
    series = [
        (profile["profileLabel"].split(" — ")[-1], [item["rate"] for item in profile["daily"]])
        for profile in analysis["routineProfiles"]
    ]
    draw_grouped_chart(
        destination,
        title="Day-over-day title overlap",
        subtitle="Share of each current first-50 rail found anywhere in the prior synthetic day",
        categories=["D1→D2", "D2→D3", "D3→D4", "D4→D5"],
        series=series,
        maximum=0.7,
    )


def build_repetition_chart(analysis: dict, destination: Path) -> None:
    series = [
        (profile["profileLabel"].split(" — ")[-1], [profile["progressive"][1]["rate"], profile["progressive"][2]["rate"]])
        for profile in analysis["repetitionProfiles"]
    ]
    draw_grouped_chart(
        destination,
        title="Progressive exact-clip recurrence",
        subtitle="Share of each 20-position run whose canonical clip appeared in an earlier run",
        categories=["Run 2", "Run 3"],
        series=series,
        maximum=0.7,
    )


def draw_grouped_chart(
    destination: Path,
    *,
    title: str,
    subtitle: str,
    categories: list[str],
    series: list[tuple[str, list[float]]],
    maximum: float,
) -> None:
    width, height = 1500, 820
    image = Image.new("RGB", (width, height), "#FCFAF4")
    draw = ImageDraw.Draw(image)
    title_font = load_font(42, bold=True)
    subtitle_font = load_font(23)
    label_font = load_font(22)
    value_font = load_font(20, bold=True)
    small_font = load_font(18)
    draw.text((92, 62), title, fill=f"#{INK}", font=title_font)
    draw.text((92, 120), subtitle, fill=f"#{MUTED}", font=subtitle_font)

    left, top, right, bottom = 128, 224, 1420, 680
    for step in range(0, 8):
        value = step / 10
        if value > maximum:
            continue
        y = bottom - (value / maximum) * (bottom - top)
        draw.line((left, y, right, y), fill="#DDD8D0", width=2)
        draw.text((52, y - 12), f"{round(value * 100):d}%", fill=f"#{MUTED}", font=small_font)

    colors = ["#96683D", "#4F779D", "#7D8B5F"]
    group_width = (right - left) / len(categories)
    bar_width = min(92, group_width / (len(series) + 1))
    for category_index, category in enumerate(categories):
        center = left + group_width * (category_index + 0.5)
        draw.text((center - 42, bottom + 28), category, fill=f"#{INK}", font=label_font)
        for series_index, (label, values) in enumerate(series):
            value = values[category_index]
            x0 = center + (series_index - (len(series) - 1) / 2) * (bar_width + 18) - bar_width / 2
            y0 = bottom - (value / maximum) * (bottom - top)
            draw.rounded_rectangle((x0, y0, x0 + bar_width, bottom), radius=8, fill=colors[series_index])
            text = f"{round(value * 100):d}%"
            draw.text((x0 + bar_width / 2 - 22, y0 - 30), text, fill=f"#{INK}", font=value_font)

    legend_x = 1020
    for index, (label, _) in enumerate(series):
        y = 70 + index * 34
        draw.rectangle((legend_x, y, legend_x + 22, y + 22), fill=colors[index])
        draw.text((legend_x + 34, y - 2), label, fill=f"#{INK}", font=small_font)

    draw.text((92, 758), "Deterministic synthetic public run", fill=f"#{MUTED}", font=small_font)
    image.save(destination, optimize=True)


def load_font(size: int, *, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def pct(value: float) -> str:
    return f"{round(value * 100):d}%"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


if __name__ == "__main__":
    main()
